import telebot
from docx import Document
from docx.shared import Mm
from docx.shared import Pt
from docx.shared import Cm



bot = telebot.TeleBot()
doc = Document('C:/Users/a.v.korolev/Desktop/Python/example.docx')

@bot.message_handler(commands=['start'])
def main(message):
    bot.send_message(message.chat.id, 'ответ на кнопку старт')

@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    photo = message.photo[-1]
    file_info = bot.get_file(photo.file_id)
    downloaded_file = bot.download_file(file_info.file_path)
    save_path = f'C:/Users/a.v.korolev/Desktop/Python/photo/{photo.file_unique_id}.jpg'
    with open(save_path, 'wb') as new_file:
        new_file.write(downloaded_file)
    table = doc.tables[0]
    cell = table.rows[1].cells[3]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(save_path, width=Mm(80))
    run = paragraph.add_run(f'\nФото {photo.file_unique_id}.jpg\n\n')
    run.font.size = Pt(8)




    doc.save('C:/Users/a.v.korolev/Desktop/Python/DocTest/.venv/example.docx')
    bot.reply_to(message, 'Фотография сохранена.')




# @bot.message_handler(content_types=['photo'])
# def handle_photo(message):
#     photo = message.photo[-1]
#     file_info = bot.get_file(photo.file_id)
#     downloaded_file = bot.download_file(file_info.file_path)
#     save_path = f'{photo.file_unique_id}.jpg'
#     with open(save_path, 'wb') as new_file:
#         new_file.write(downloaded_file)
#     doc.add_picture(f'C:/Users/a.v.korolev/Desktop/Python/{save_path}', width=Mm(50))
#     doc.save('example.docx')
#     bot.reply_to(message, 'Фотография сохранена.')
#
#
# @bot.message_handler(content_types=['document', 'video', 'audio', 'voice', 'sticker'])
# def handle_file(message):
#     file_info = bot.get_file(message.document.file_id)
#     downloaded_file = bot.download_file(file_info.file_path)
#     save_path = message.document.file_name  # сохраняем файл с его исходным именем
#
#     with open(save_path, 'wb') as new_file:
#         new_file.write(downloaded_file)
#     doc.add_picture(f'C:/Users/a.v.korolev/Desktop/Python/{save_path}', width=Mm(50))
#     doc.save('example.docx')
#
#     bot.reply_to(message, 'Файл сохранен.')


bot.polling()
